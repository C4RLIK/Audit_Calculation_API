from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import JSONResponse, HTMLResponse, FileResponse
from pydantic import BaseModel, Field
from typing import List, Optional
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import secrets
import time
from fastapi.templating import Jinja2Templates
import os
from io import BytesIO
from fastapi.openapi.utils import get_openapi

app = FastAPI()
templates = Jinja2Templates(directory="templates")

# Хранилище сессий
form_sessions = {}

class Indicator(BaseModel):
    """Финансовый показатель для расчёта существенности"""
    name: str = Field(..., example="Выручка от продаж", description="Название финансового показателя")
    value: float = Field(..., example=1800000, description="Значение показателя в рублях")

class CalculationRequest(BaseModel):
    """Запрос на расчёт уровня существенности"""
    indicators: List[Indicator] = Field(..., example=[
        {"name": "Выручка от продаж", "value": 1800000},
        {"name": "Себестоимость продаж", "value": 1374000},
        {"name": "Прибыль от продаж", "value": 480000},
        {"name": "Чистая прибыль", "value": 480000},
        {"name": "Чистая прибыль (повтор)", "value": 668000},
        {"name": "Уставный капитал", "value": 100000},
        {"name": "Основные средства", "value": 208000}
    ], description="Список финансовых показателей")
    deviation_threshold: float = Field(50, example=50, description="Порог отклонения в процентах (0-100)", ge=0, le=100)
    rounding_limit: float = Field(50, example=50, description="Максимальное отклонение при округлении", ge=0)
    with_docx: bool = Field(False, example=False, description="Генерировать Word-документ с отчётом")

class DeviationInfo(BaseModel):
    """Информация об отклонении показателя"""
    absolute: float = Field(..., example=-142857.14, description="Абсолютное отклонение от среднего")
    percent: float = Field(..., example=-16.67, description="Отклонение в процентах от среднего")

class IndicatorResult(BaseModel):
    """Результат расчёта для одного показателя"""
    name: str = Field(..., example="Выручка от продаж", description="Название показателя")
    value: float = Field(..., example=1800000, description="Значение показателя")
    deviation: DeviationInfo = Field(..., description="Информация об отклонении")

class CalculationSteps(BaseModel):
    """Детали расчёта уровня существенности"""
    initial_mean: float = Field(..., example=857142.86, description="Первоначальное среднее значение")
    filtered_mean: float = Field(..., example=857142.86, description="Среднее после исключения выбросов")
    excluded_count: int = Field(..., example=0, description="Количество исключённых показателей")
    excluded_values: List[float] = Field(..., example=[], description="Значения исключённых показателей")
    indicators: List[IndicatorResult] = Field(..., description="Результаты по всем показателям")
    rounded_value: float = Field(..., example=857100.0, description="Округлённое значение")

class CalculationResponse(BaseModel):
    """Ответ с результатом расчёта уровня существенности"""
    materiality_level: float = Field(..., example=857100.0, description="Итоговый уровень существенности")
    calculation_steps: CalculationSteps = Field(..., description="Детализированные шаги расчёта")
    indicators: List[Indicator] = Field(..., description="Исходные показатели")
    message: str = Field("Расчёт выполнен успешно", example="Расчёт выполнен успешно", description="Статус выполнения")

# Функция расчёта (без изменений)
def calculate_materiality(data, deviation_threshold, rounding_limit):
    try:
        values = [indicator.value for indicator in data]
        if len(values) == 0:
            return None, "Нет данных для расчёта"
        
        mean = np.mean(values)
        deviations = [(x, abs(x - mean) / mean * 100) for x in values]
        
        filtered = [x for x, dev in deviations if dev <= deviation_threshold]
        excluded = [x for x, dev in deviations if dev > deviation_threshold]
        
        if not filtered:
            return None, "Все показатели исключены как нерепрезентативные"
        
        new_mean = np.mean(filtered)
        rounded = round(new_mean / 100) * 100
        if abs(rounded - new_mean) > rounding_limit:
            rounded = new_mean
        
        details = {
            "initial_mean": mean,
            "deviations": deviations,
            "excluded": excluded,
            "filtered": filtered,
            "new_mean": new_mean,
            "rounded": rounded,
            "indicator_names": [indicator.name for indicator in data]
        }
        
        return rounded, details
    
    except Exception as e:
        return None, f"Ошибка расчёта: {str(e)}"

# Улучшенная генерация Word-документа
def create_word_report(details, deviation_threshold, indicators):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Заголовок
    title = doc.add_heading('Расчёт уровня существенности', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Исходные данные
    doc.add_heading('1. Исходные данные:', level=2)
    for idx, (indicator, value) in enumerate(zip(indicators, [x for x, _ in details["deviations"]]), 1):
        doc.add_paragraph(f"{idx}. {indicator.name}: {value:,.0f} руб.", style='ListNumber')

    # 2. Среднее арифметическое
    doc.add_heading('2. Расчёт среднего арифметического:', level=2)
    values_str = " + ".join([f"{x:,.0f}" for x, _ in details["deviations"]])
    doc.add_paragraph(f"({values_str}) / {len(details['deviations'])} = {details['initial_mean']:,.0f} руб.")

    # 3. Отклонения показателей
    doc.add_heading('3. Определение отклонений показателей от среднего:', level=2)
    for x, dev in details["deviations"]:
        deviation = (x - details['initial_mean'])/details['initial_mean']*100
        doc.add_paragraph(f"• Отклонение: {deviation:+.2f}% от среднего", style='ListBullet')

    # 4. Исключение показателей
    doc.add_heading(f'4. Исключение показателей с отклонением > {deviation_threshold}%:', level=2)
    if details["excluded"]:
        for x in details["excluded"]:
            doc.add_paragraph(f"• Исключён показатель: {x:,.0f} руб.", style='ListBullet')
    else:
        doc.add_paragraph("Нет исключённых показателей")

    # 5. Новое среднее
    doc.add_heading('5. Расчёт нового среднего арифметического:', level=2)
    doc.add_paragraph(f"({' + '.join([f'{x:,.0f}' for x in details['filtered']])}) / {len(details['filtered'])} = {details['new_mean']:,.2f} руб.")

    # 6. Округление
    doc.add_heading('6. Округление результата:', level=2)
    doc.add_paragraph(f"Округлённое значение: {details['rounded']:,.0f} руб.")

    # 7. Итог
    doc.add_heading('7. Итоговый уровень существенности:', level=2)
    p = doc.add_paragraph()
    p.add_run(f"{details['rounded']:,.0f} рублей").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc

# Улучшенный JSON-ответ
def format_response(result, details, indicators):
    formatted_deviations = []
    for (value, percent), name in zip(details["deviations"], details["indicator_names"]):
        deviation_value = value - details["initial_mean"]
        deviation_percent = (value - details["initial_mean"]) / details["initial_mean"] * 100
        formatted_deviations.append(IndicatorResult(
            name=name,
            value=value,
            deviation=DeviationInfo(
                absolute=deviation_value,
                percent=deviation_percent
            )
        ))

    return CalculationResponse(
        materiality_level=result,
        calculation_steps=CalculationSteps(
            initial_mean=details["initial_mean"],
            filtered_mean=details["new_mean"],
            excluded_count=len(details["excluded"]),
            excluded_values=details["excluded"],
            indicators=formatted_deviations,
            rounded_value=details["rounded"]
        ),
        indicators=indicators,
        message="Расчёт выполнен успешно"
    )

@app.post("/api/v1/calculate", response_model=CalculationResponse)
async def calculate_materiality_endpoint(request: CalculationRequest):
    """
    Расчёт уровня существенности на основе финансовых показателей
    
    Пример запроса:
    ```json
    {
        "indicators": [
            {"name": "Выручка от продаж", "value": 1800000},
            {"name": "Себестоимость продаж", "value": 1374000},
            {"name": "Прибыль от продаж", "value": 480000},
            {"name": "Чистая прибыль", "value": 480000},
            {"name": "Чистая прибыль (повтор)", "value": 668000},
            {"name": "Уставный капитал", "value": 100000},
            {"name": "Основные средства", "value": 208000}
        ],
        "deviation_threshold": 50,
        "rounding_limit": 50,
        "with_docx": false
    }
    ```
    """
    if len(request.indicators) > 50:
        raise HTTPException(status_code=400, detail="Максимум 50 показателей")

    result, details = calculate_materiality(request.indicators, request.deviation_threshold, request.rounding_limit)
    
    if result is None:
        raise HTTPException(status_code=400, detail=details)

    if not request.with_docx:
        return format_response(result, details, request.indicators)

    doc = create_word_report(details, request.deviation_threshold, request.indicators)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return FileResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="materiality_report.docx"
    )

# Остальные эндпоинты без изменений
@app.get("/api/v1/generate-form")
async def generate_form_session(request: Request):
    session_id = secrets.token_urlsafe(16)
    expires_at = time.time() + 600  # 10 минут
    
    form_sessions[session_id] = {
        "expires_at": expires_at,
        "used": False
    }
    
    # Возвращаем полный URL
    form_url = str(request.url_for("calculation_form", session_id=session_id))
    return {"form_url": form_url, "expires_at": expires_at}

@app.get("/form/{session_id}", response_class=HTMLResponse)
async def calculation_form(request: Request, session_id: str):
    session = form_sessions.get(session_id)
    
    if not session:
        raise HTTPException(status_code=404, detail="Сессия не найдена")
    
    if time.time() > session["expires_at"]:
        raise HTTPException(status_code=410, detail="Время действия сессии истекло")
    
    if session["used"]:
        raise HTTPException(status_code=403, detail="Форма уже была использована")
    
    form_sessions[session_id]["used"] = True
    
    return templates.TemplateResponse("form.html", {
        "request": request,
        "session_id": session_id,
        "expires_at": session["expires_at"]
    })

# Кастомизация документации OpenAPI
def custom_openapi():
    if app.openapi_schema:
        return app.openapi_schema
    
    openapi_schema = get_openapi(
        title="API для расчёта уровня существенности",
        version="1.0.0",
        description="API для автоматического расчёта уровня существенности на основе финансовых показателей",
        routes=app.routes,
    )
    
    # Добавляем примеры для схем
    openapi_schema["components"]["schemas"]["CalculationRequest"]["example"] = {
        "indicators": [
            {"name": "Выручка от продаж", "value": 1800000},
            {"name": "Себестоимость продаж", "value": 1374000},
            {"name": "Прибыль от продаж", "value": 480000},
            {"name": "Чистая прибыль", "value": 480000},
            {"name": "Чистая прибыль (повтор)", "value": 668000},
            {"name": "Уставный капитал", "value": 100000},
            {"name": "Основные средства", "value": 208000}
        ],
        "deviation_threshold": 50,
        "rounding_limit": 50,
        "with_docx": False
    }
    
    app.openapi_schema = openapi_schema
    return app.openapi_schema

app.openapi = custom_openapi

# Создаем директорию для шаблонов, если её нет
os.makedirs("templates", exist_ok=True)

# Создаем шаблон формы (без изменений)
with open("templates/form.html", "w", encoding="utf-8") as f:
    f.write("""
<!DOCTYPE html>
<html>
<head>
    <title>Расчёт уровня существенности</title>
    <style>
        body { font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }
        .timer { color: red; font-weight: bold; }
        form { margin-top: 20px; }
        .indicator { margin-bottom: 15px; display: flex; gap: 10px; align-items: center; }
        button { margin-top: 20px; padding: 10px 20px; }
        .result { margin-top: 20px; padding: 15px; background: #f5f5f5; border-radius: 5px; }
    </style>
</head>
<body>
    <h1>Калькулятор уровня существенности</h1>
    <p>Осталось времени: <span class="timer" id="timer">10:00</span></p>
    
    <form id="calculationForm">
        <div id="indicatorsContainer">
            <div class="indicator">
                <label>Показатель 1:</label>
                <input type="text" name="name_1" placeholder="Название" required>
                <input type="number" name="value_1" placeholder="Значение" min="0" step="1000" required>
            </div>
        </div>
        
        <button type="button" onclick="addIndicator()">+ Добавить показатель</button>
        
        <div style="margin-top: 20px;">
            <label>Допустимое отклонение (%):</label>
            <input type="number" name="deviation_threshold" value="50" min="0" max="100" required>
        </div>
        
        <div>
            <label>Макс. отклонение при округлении:</label>
            <input type="number" name="rounding_limit" value="50" min="0" required>
        </div>
        
        <div>
            <label>
                <input type="checkbox" name="with_docx">
                Сгенерировать Word-отчёт
            </label>
        </div>
        
        <button type="submit">Рассчитать</button>
    </form>
    
    <div id="result" class="result"></div>
    
    <script>
        // Таймер
        const expiresAt = {{ expires_at }};
        function updateTimer() {
            const now = Math.floor(Date.now() / 1000);
            const remaining = expiresAt - now;
            
            if (remaining <= 0) {
                document.getElementById('timer').textContent = "00:00";
                alert("Время сессии истекло!");
                return;
            }
            
            const minutes = Math.floor(remaining / 60);
            const seconds = remaining % 60;
            document.getElementById('timer').textContent = 
                `${minutes.toString().padStart(2, '0')}:${seconds.toString().padStart(2, '0')}`;
            
            setTimeout(updateTimer, 1000);
        }
        
        updateTimer();
        
        // Добавление показателей
        let indicatorCount = 1;
        function addIndicator() {
            indicatorCount++;
            if (indicatorCount > 50) {
                alert("Максимум 50 показателей");
                return;
            }
            
            const container = document.getElementById('indicatorsContainer');
            const newIndicator = document.createElement('div');
            newIndicator.className = 'indicator';
            newIndicator.innerHTML = `
                <label>Показатель ${indicatorCount}:</label>
                <input type="text" name="name_${indicatorCount}" placeholder="Название" required>
                <input type="number" name="value_${indicatorCount}" placeholder="Значение" min="0" step="1000" required>
            `;
            container.appendChild(newIndicator);
        }
        
        // Отправка формы
        document.getElementById('calculationForm').addEventListener('submit', async function(e) {
            e.preventDefault();
            
            const formData = new FormData(this);
            const indicators = [];
            
            for (let i = 1; i <= indicatorCount; i++) {
                const name = formData.get(`name_${i}`);
                const value = parseFloat(formData.get(`value_${i}`));
                
                if (name && !isNaN(value)) {
                    indicators.push({ name, value });
                }
            }
            
            const requestData = {
                indicators,
                deviation_threshold: parseFloat(formData.get('deviation_threshold')),
                rounding_limit: parseFloat(formData.get('rounding_limit')),
                with_docx: formData.get('with_docx') === 'on'
            };
            
            try {
                const response = await fetch('/api/v1/calculate', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify(requestData)
                });
                
                if (requestData.with_docx && response.ok) {
                    // Скачивание Word-документа
                    const blob = await response.blob();
                    const url = window.URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = 'materiality_report.docx';
                    document.body.appendChild(a);
                    a.click();
                    document.body.removeChild(a);
                } else {
                    // Отображение JSON-результата
                    const result = await response.json();
                    displayResult(result);
                }
            } catch (error) {
                document.getElementById('result').innerHTML = 
                    `<p style="color: red;">Ошибка: ${error.message}</p>`;
            }
        });
        
        function displayResult(result) {
            let html = '<h2>Результаты расчёта</h2>';
            
            html += '<h3>Исходные показатели:</h3><ul>';
            result.indicators.forEach(ind => {
                html += `<li>${ind.name}: ${ind.value.toLocaleString()} руб.</li>`;
            });
            html += '</ul>';
            
            html += '<h3>Процесс расчёта:</h3>';
            html += `<p>Среднее арифметическое: <b>${result.calculation_steps.initial_mean.toLocaleString()} руб.</b></p>`;
            
            if (result.calculation_steps.excluded_count > 0) {
                html += `<p>Исключено показателей: ${result.calculation_steps.excluded_count}</p>`;
                html += '<p>Исключённые значения: ' + 
                    result.calculation_steps.excluded_values.map(v => v.toLocaleString() + ' руб.').join(', ') + '</p>';
            }
            
            html += `<p>Среднее после исключения: <b>${result.calculation_steps.filtered_mean.toLocaleString()} руб.</b></p>`;
            html += `<p>Округлённое значение: <b>${result.calculation_steps.rounded_value.toLocaleString()} руб.</b></p>`;
            
            html += '<h3>Отклонения показателей:</h3><ul>';
            result.calculation_steps.indicators.forEach(ind => {
                html += `<li>${ind.name}: ${ind.value.toLocaleString()} руб. ` +
                        `(отклонение: ${ind.deviation.percent.toFixed(2)}%)</li>`;
            });
            html += '</ul>';
            
            html += `<h2 style="color: green;">Итоговый уровень существенности: ${result.materiality_level.toLocaleString()} руб.</h2>`;
            
            document.getElementById('result').innerHTML = html;
        }
    </script>
</body>
</html>
""")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
